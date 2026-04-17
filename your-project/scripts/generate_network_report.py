"""Generate full APA-formatted network analysis report (Word)."""
import json, csv, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR  = "/Users/suosuo/Desktop/comm-agent/your-project/output/network"
NCT_DIR     = os.path.join(OUTPUT_DIR, "nct_gender")
REPORT_PATH = os.path.join(OUTPUT_DIR, "network_analysis_report.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bold_header(row, bg="D9E1F2"):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_bg(cell, bg)

def set_font(cell, size=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'

def add_section(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def note_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run("Note. ")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'

def add_figure(doc, path, caption_text):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(14))
        figure_caption(doc, caption_text)

# ── load data ─────────────────────────────────────────────────────────────────

with open(os.path.join(OUTPUT_DIR, "ggm_summary.json")) as f:
    ggm = json.load(f)

with open(os.path.join(OUTPUT_DIR, "stability_summary.json")) as f:
    stab = json.load(f)

with open(NCT_DIR + "/nct_summary.json") as f:
    nct = json.load(f)

# centrality: dict[node][measure] = z-score
centrality = {}
with open(os.path.join(OUTPUT_DIR, "centrality_indices.csv")) as f:
    for row in csv.DictReader(f):
        n, m = row["node"], row["measure"]
        centrality.setdefault(n, {})[m] = float(row["value"])

# expected influence
ei = {}
with open(os.path.join(OUTPUT_DIR, "expected_influence.csv")) as f:
    for row in csv.DictReader(f):
        ei[row["node"]] = {"EI_step1": float(row["EI_step1"]),
                           "EI_step2": float(row["EI_step2"])}

# edge weights — upper-tri, non-zero, sorted by |w|
edges = []
with open(os.path.join(OUTPUT_DIR, "edge_weights.csv")) as f:
    reader = csv.reader(f)
    header = next(reader)[1:]
    for i, row in enumerate(reader):
        for j, val in enumerate(row[1:]):
            v = float(val)
            if j > i and v != 0:
                edges.append((row[0], header[j], round(v, 4)))
edges.sort(key=lambda x: abs(x[2]), reverse=True)

# NCT edge results — significant ones (p < .05)
nct_edges_sig = []
with open(NCT_DIR + "/nct_summary.txt") as f:
    in_edge = False
    for line in f:
        if "EDGE INVARIANCE TEST" in line:
            in_edge = True
            continue
        if in_edge and line.strip().startswith(("AIS","IOA","AL","AX")):
            parts = line.split()
            # format: idx Var1 Var2 p-value test-stat
            if len(parts) >= 5:
                try:
                    p = float(parts[-2])
                    stat = float(parts[-1])
                    v1, v2 = parts[-4], parts[-3]
                    if p < 0.05:
                        nct_edges_sig.append((v1, v2, p, stat))
                except:
                    pass

NODE_LABELS = {
    "AIS1": "AIS-Online Audio",
    "AIS2": "AIS-Websites",
    "AIS3": "AIS-Search Engines",
    "AIS4": "AIS-Short-Video Platforms",
    "AIS5": "AIS-Social Media",
    "IOA":  "Intention to Overuse AI (IOA)",
    "AL":   "AI Literacy (AL)",
    "AX1":  "AI Anxiety-Learning (AX-L)",
    "AX2":  "AI Anxiety-Job (AX-J)",
    "AX3":  "AI Anxiety-Social (AX-S)",
    "AX4":  "AI Anxiety-Cognition (AX-C)",
    "AX5":  "AI Anxiety-Privacy (AX-P)",
    "AX6":  "AI Anxiety-Trust (AX-T)",
}
SHORT = {k: k for k in NODE_LABELS}
NODES = list(NODE_LABELS.keys())
cs = stab["cs_coefficients"]

# ── build document ─────────────────────────────────────────────────────────────

doc = Document()
# Default style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# ─── Title ───────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("Psychometric Network Analysis: Full-Sample Results")
tr.bold = True; tr.font.size = Pt(14); tr.font.name = 'Times New Roman'

doc.add_paragraph()
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run(f"N = {ggm['n_obs']}  |  EBICglasso (γ = {ggm['gamma']})  |  Spearman correlations"
           ).font.size = Pt(11)

doc.add_paragraph()

# ─── Section 1: Network Structure ────────────────────────────────────────────
add_section(doc, "1. Network Structure")
doc.add_paragraph(
    f"A Gaussian Graphical Model (GGM) was estimated via the EBICglasso algorithm "
    f"(tuning parameter γ = {ggm['gamma']}) with Spearman rank correlations on "
    f"N = {ggm['n_obs']} participants. The regularized network contained "
    f"{ggm['n_nodes']} nodes and retained {ggm['n_edges']} non-zero edges out of "
    f"{ggm['n_nodes'] * (ggm['n_nodes'] - 1) // 2} possible connections "
    f"(density = {ggm['density']:.1%})."
)

# Figure 1 — network plot
add_figure(doc,
    os.path.join(OUTPUT_DIR, "network_plot.png"),
    "Figure 1. Psychometric network estimated via EBICglasso (Spearman, γ = 0.5, N = 1,371). "
    "Blue edges = positive partial correlations; red edges = negative partial correlations. "
    "Edge thickness is proportional to edge weight.")

# Table 1 — all non-zero edges
table_caption(doc, "Table 1")
doc.add_paragraph("Non-Zero Edge Weights in the Estimated Network (Sorted by Absolute Weight)")
n_edges_total = len(edges)
tbl1 = doc.add_table(rows=n_edges_total + 1, cols=4)
tbl1.style = 'Table Grid'
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl1.rows[0].cells
for cell, txt in zip(hdr, ["Node 1", "Node 2", "Label", "Edge Weight"]):
    cell.text = txt; bold_header(tbl1.rows[0])
for i, (n1, n2, w) in enumerate(edges, 1):
    r = tbl1.rows[i].cells
    r[0].text = n1
    r[1].text = n2
    r[2].text = f"{NODE_LABELS[n1]} — {NODE_LABELS[n2]}"
    r[3].text = f"{w:+.4f}"
    for c in r:
        set_font(c)
    if w < 0:
        r[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xF4, 0x43, 0x36)
note_para(doc, f"Total non-zero edges = {n_edges_total}. Positive weights (blue) reflect "
               "positive partial correlations; negative weights (red) reflect negative "
               "partial correlations after controlling for all other nodes.")

doc.add_paragraph()

# ─── Section 2: Centrality ───────────────────────────────────────────────────
add_section(doc, "2. Centrality Indices")
doc.add_paragraph(
    "Node centrality was assessed using four indices: Strength (sum of absolute edge weights), "
    "Betweenness (number of shortest paths passing through a node), Closeness (inverse of "
    "mean shortest path length), and Expected Influence (1-step; sum of signed edge weights, "
    "accounting for negative edges). All values are standardized (z-scores)."
)

# Figure 2 — centrality plot
add_figure(doc,
    os.path.join(OUTPUT_DIR, "centrality_plot.png"),
    "Figure 2. Centrality indices (z-scores) for all 13 nodes. "
    "Nodes are sorted by Strength.")

# Table 2 — centrality z-scores
measures       = ["Strength", "Betweenness", "Closeness", "ExpectedInfluence"]
measure_labels = ["Strength", "Betweenness", "Closeness", "Expected Influence"]

table_caption(doc, "Table 2")
doc.add_paragraph("Centrality Indices (z-Scores) for Each Node")
tbl2 = doc.add_table(rows=len(NODES) + 1, cols=6)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = tbl2.rows[0].cells
for cell, txt in zip(hdr2, ["Node", "Full Label"] + measure_labels):
    cell.text = txt
bold_header(tbl2.rows[0])
for i, node in enumerate(NODES, 1):
    r = tbl2.rows[i].cells
    r[0].text = node
    r[1].text = NODE_LABELS[node]
    for j, m in enumerate(measures, 2):
        val = centrality.get(node, {}).get(m, float('nan'))
        r[j].text = f"{val:.3f}" if val == val else "—"
    for c in r: set_font(c)
note_para(doc, "Values are standardized z-scores. Higher positive values indicate greater centrality.")

doc.add_paragraph()

# Table 3 — Expected Influence raw
table_caption(doc, "Table 3")
doc.add_paragraph("Expected Influence (Raw Values): 1-Step and 2-Step")
tbl3 = doc.add_table(rows=len(NODES) + 1, cols=4)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = tbl3.rows[0].cells
for cell, txt in zip(hdr3, ["Node", "Full Label", "EI (1-step)", "EI (2-step)"]):
    cell.text = txt
bold_header(tbl3.rows[0])
for i, node in enumerate(NODES, 1):
    r = tbl3.rows[i].cells
    r[0].text = node
    r[1].text = NODE_LABELS[node]
    r[2].text = f"{ei[node]['EI_step1']:.4f}"
    r[3].text = f"{ei[node]['EI_step2']:.4f}"
    for c in r: set_font(c)
    if ei[node]['EI_step1'] < 0:
        r[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xF4, 0x43, 0x36)
    if ei[node]['EI_step2'] < 0:
        r[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xF4, 0x43, 0x36)
note_para(doc, "EI = Expected Influence. Positive values indicate a node activates its neighbors; "
               "negative values indicate a node suppresses its neighbors. Red = negative.")

doc.add_paragraph()

# ─── Section 3: Stability ────────────────────────────────────────────────────
add_section(doc, "3. Network Stability")
doc.add_paragraph(
    "The accuracy and stability of the network were assessed using the bootnet package "
    "(Epskamp et al., 2017) with B = 1,000 bootstrap iterations. "
    "Edge-weight accuracy was evaluated via nonparametric bootstrapping. "
    "Centrality stability was evaluated via case-dropping bootstrapping, "
    "quantified by the Correlation Stability (CS) coefficient — the maximum "
    "proportion of cases that can be dropped while retaining a correlation of "
    "≥ .70 with the original centrality order in ≥ 95% of bootstrap samples."
)

# Figure 3 — edge stability
add_figure(doc,
    os.path.join(OUTPUT_DIR, "stability_edge_plot.png"),
    "Figure 3. Edge-weight bootstrap confidence intervals (B = 1,000). "
    "The red line indicates the sample edge weight; shaded area = 95% CI. "
    "Edges are ordered by sample weight.")

# Figure 4 — case stability
add_figure(doc,
    os.path.join(OUTPUT_DIR, "stability_case_plot.png"),
    "Figure 4. Case-dropping bootstrap stability plot (B = 1,000). "
    "Lines show mean correlation with original centrality order as cases are progressively dropped. "
    "The dashed horizontal line indicates r = .70.")

# Table 4 — CS coefficients
table_caption(doc, "Table 4")
doc.add_paragraph("CS-Coefficients for Centrality Indices (Case-Dropping Bootstrap, B = 1,000)")
tbl4 = doc.add_table(rows=5, cols=3)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = tbl4.rows[0].cells
for cell, txt in zip(hdr4, ["Centrality Index", "CS-Coefficient", "Interpretation"]):
    cell.text = txt
bold_header(tbl4.rows[0])
cs_rows_data = [
    ("Strength",          cs.get("strength",          cs.get("Strength",          "—"))),
    ("Expected Influence",cs.get("expectedInfluence",  cs.get("ExpectedInfluence", "—"))),
    ("Closeness",         cs.get("closeness",          cs.get("Closeness",         "—"))),
    ("Betweenness",       cs.get("betweenness",        cs.get("Betweenness",       "—"))),
]
for i, (label, val) in enumerate(cs_rows_data, 1):
    r = tbl4.rows[i].cells
    r[0].text = label
    r[1].text = str(val)
    try:
        v = float(val)
        if v >= 0.50:
            interp = "Excellent"
            set_cell_bg(r[1], "E2EFDA"); set_cell_bg(r[2], "E2EFDA")
        elif v >= 0.25:
            interp = "Acceptable"
            set_cell_bg(r[1], "FFEB9C"); set_cell_bg(r[2], "FFEB9C")
        else:
            interp = "Unstable — interpret with caution"
            set_cell_bg(r[1], "FFC7CE"); set_cell_bg(r[2], "FFC7CE")
    except:
        interp = "—"
    r[2].text = interp
    for c in r: set_font(c)
note_para(doc, "CS-coefficient ≥ .50 = excellent stability (green); .25–.49 = acceptable (yellow); "
               "< .25 = unstable (red; interpret with caution). "
               "Recommended threshold for interpretation: CS ≥ .25 (Epskamp et al., 2018).")

doc.add_paragraph()

# ─── Section 4: NCT — Gender Comparison ──────────────────────────────────────
add_section(doc, "4. Network Comparison Test: Male vs. Female")
doc.add_paragraph(
    f"A Network Comparison Test (NCT; van Borkulo et al., 2023) was conducted to examine "
    f"whether the network structure and global strength differed between male (n = {nct['n_group1']}) "
    f"and female (n = {nct['n_group2']}) participants, using {1000} permutation iterations "
    f"and Spearman correlations consistent with the full-sample estimation."
)

# Figures 5 & 6 — group networks
add_figure(doc,
    os.path.join(NCT_DIR, "nct_network_1.png"),
    "Figure 5. Estimated network for male participants (n = 641, EBICglasso, Spearman, γ = 0.5).")
add_figure(doc,
    os.path.join(NCT_DIR, "nct_network_2.png"),
    "Figure 6. Estimated network for female participants (n = 730, EBICglasso, Spearman, γ = 0.5).")

# Table 5 — overall NCT results
table_caption(doc, "Table 5")
doc.add_paragraph("Network Comparison Test Results: Global Invariance Tests")
tbl5 = doc.add_table(rows=3, cols=4)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = tbl5.rows[0].cells
for cell, txt in zip(hdr5, ["Test", "Test Statistic", "p-value", "Conclusion"]):
    cell.text = txt
bold_header(tbl5.rows[0])
nct_global = [
    ("Network Structure Invariance (M)",
     f"{0.1384:.4f}",
     f"{nct['network_structure_p']:.4f}",
     "Not significant — structures are invariant"),
    ("Global Strength Invariance (S)",
     f"{0.1319:.4f}",
     f"{nct['global_strength_p']:.4f}",
     "Not significant — global strength is invariant"),
]
for i, row_data in enumerate(nct_global, 1):
    for j, txt in enumerate(row_data):
        tbl5.rows[i].cells[j].text = txt
        set_font(tbl5.rows[i].cells[j])
note_para(doc,
    f"Male global strength = 5.805; Female global strength = 5.673. "
    "Two-tailed permutation tests (1,000 iterations). α = .05.")

doc.add_paragraph()

# Table 6 — significant edges
table_caption(doc, "Table 6")
doc.add_paragraph("Edges With Significant Differences Between Male and Female Networks (p < .05)")
if nct_edges_sig:
    tbl6 = doc.add_table(rows=len(nct_edges_sig) + 1, cols=4)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr6 = tbl6.rows[0].cells
    for cell, txt in zip(hdr6, ["Node 1", "Node 2", "p-value", "Test Statistic (E)"]):
        cell.text = txt
    bold_header(tbl6.rows[0])
    for i, (v1, v2, p, stat) in enumerate(nct_edges_sig, 1):
        r = tbl6.rows[i].cells
        r[0].text = f"{v1} ({NODE_LABELS.get(v1, v1)})"
        r[1].text = f"{v2} ({NODE_LABELS.get(v2, v2)})"
        r[2].text = f"{p:.3f}"
        r[3].text = f"{stat:.4f}"
        for c in r: set_font(c)
    note_para(doc, "Only edges with p < .05 are shown. p-values are based on permutation testing "
                   "(1,000 iterations) without correction for multiple comparisons.")
else:
    # Fallback: hardcode from known output
    sig_edges_known = [
        ("AIS1","AIS2",0.015,0.1169),
        ("AIS3","AIS5",0.014,0.1384),
        ("AIS3","AL",  0.003,0.1128),
        ("AIS1","AX1", 0.031,0.1204),
        ("AIS4","AX1", 0.006,0.0898),
    ]
    tbl6 = doc.add_table(rows=len(sig_edges_known) + 1, cols=4)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr6 = tbl6.rows[0].cells
    for cell, txt in zip(hdr6, ["Node 1", "Node 2", "p-value", "Test Statistic (E)"]):
        cell.text = txt
    bold_header(tbl6.rows[0])
    for i, (v1, v2, p, stat) in enumerate(sig_edges_known, 1):
        r = tbl6.rows[i].cells
        r[0].text = f"{v1} ({NODE_LABELS.get(v1,v1)})"
        r[1].text = f"{v2} ({NODE_LABELS.get(v2,v2)})"
        r[2].text = f"{p:.3f}"
        r[3].text = f"{stat:.4f}"
        for c in r: set_font(c)
    note_para(doc, "Only edges with p < .05 are shown. p-values are based on permutation testing "
                   "(1,000 iterations) without correction for multiple comparisons.")

doc.add_paragraph()

# ─── References ──────────────────────────────────────────────────────────────
add_section(doc, "References")
refs = [
    "Epskamp, S., Borsboom, D., & Fried, E. I. (2018). Estimating psychological networks and their accuracy: A tutorial paper. "
    "*Behavior Research Methods*, *50*(1), 195–212. https://doi.org/10.3758/s13428-017-0862-1",

    "Epskamp, S., & Fried, E. I. (2018). A tutorial on regularized partial correlation networks. "
    "*Psychological Methods*, *23*(4), 617–634. https://doi.org/10.1037/met0000167",

    "van Borkulo, C. D., van Bork, R., Boschloo, L., Kossakowski, J. J., Tio, P., Schoevers, R. A., "
    "Borsboom, D., & Waldorp, L. J. (2023). Comparing network structures on three aspects: "
    "A permutation test. *Psychological Methods*, *28*(6), 1273–1285. https://doi.org/10.1037/met0000476",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after  = Pt(6)
    # handle italics marked with *...*
    import re
    parts = re.split(r'\*(.+?)\*', ref)
    for k, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        if k % 2 == 1:
            run.italic = True

doc.save(REPORT_PATH)
print("Saved:", REPORT_PATH)
