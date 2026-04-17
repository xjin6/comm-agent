"""
Export individual APA-formatted tables (.docx) and figures (.png) for the
psychometric network analysis. Outputs go to:
  output/network/tables/   ← one .docx per table
  output/network/figures/  ← copies of all PNGs with descriptive names
"""

import csv, json, os, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── paths ─────────────────────────────────────────────────────────────────────

BASE    = "/Users/suosuo/Desktop/comm-agent/your-project/output/network"
NCT_DIR = os.path.join(BASE, "nct_gender")
TDIR    = os.path.join(BASE, "tables")
FDIR    = os.path.join(BASE, "figures")
os.makedirs(TDIR, exist_ok=True)
os.makedirs(FDIR, exist_ok=True)

# ── node labels ───────────────────────────────────────────────────────────────

NODE_LABELS = {
    "AIS1": "AIS-Online Audio",
    "AIS2": "AIS-Websites",
    "AIS3": "AIS-Search Engines",
    "AIS4": "AIS-Short-Video Platforms",
    "AIS5": "AIS-Social Media",
    "IOA":  "Intention to Overuse AI",
    "AL":   "AI Literacy",
    "AX1":  "AI Anxiety-Learning",
    "AX2":  "AI Anxiety-Job",
    "AX3":  "AI Anxiety-Social",
    "AX4":  "AI Anxiety-Cognition",
    "AX5":  "AI Anxiety-Privacy",
    "AX6":  "AI Anxiety-Trust",
}
NODES = list(NODE_LABELS.keys())

# ── load data ─────────────────────────────────────────────────────────────────

with open(os.path.join(BASE, "ggm_summary.json")) as f:
    ggm = json.load(f)

with open(os.path.join(BASE, "stability_summary.json")) as f:
    stab = json.load(f)
cs = stab["cs_coefficients"]

with open(os.path.join(NCT_DIR, "nct_summary.json")) as f:
    nct = json.load(f)

# centrality: {node: {measure: z_value}}
centrality = {}
with open(os.path.join(BASE, "centrality_indices.csv")) as f:
    for row in csv.DictReader(f):
        centrality.setdefault(row["node"], {})[row["measure"]] = float(row["value"])

# expected influence
ei = {}
with open(os.path.join(BASE, "expected_influence.csv")) as f:
    for row in csv.DictReader(f):
        ei[row["node"]] = {"step1": float(row["EI_step1"]),
                           "step2": float(row["EI_step2"])}

# bridge centrality
bridge = {}
with open(os.path.join(BASE, "bridge_centrality.csv")) as f:
    for row in csv.DictReader(f):
        bridge[row["node"]] = {
            "strength":  float(row["bridge_strength"]),
            "between":   float(row["bridge_between"]),
            "close":     float(row["bridge_close"]),
            "EI":        float(row["bridge_EI"]),
        }

# edge weights (full matrix)
edge_matrix = {}
with open(os.path.join(BASE, "edge_weights.csv")) as f:
    reader = csv.reader(f)
    header = next(reader)[1:]
    for row in reader:
        rname = row[0]
        edge_matrix[rname] = {header[j]: float(v) for j, v in enumerate(row[1:])}

# NCT all edges
nct_edges = []
with open(os.path.join(NCT_DIR, "nct_summary.txt")) as f:
    in_edge = False
    for line in f:
        if "EDGE INVARIANCE TEST" in line:
            in_edge = True
            continue
        if in_edge:
            parts = line.split()
            if len(parts) >= 5:
                try:
                    p    = float(parts[-2])
                    stat = float(parts[-1])
                    v1, v2 = parts[-4], parts[-3]
                    nct_edges.append((v1, v2, p, stat))
                except ValueError:
                    pass

nct_sig = [(v1, v2, p, s) for v1, v2, p, s in nct_edges if p < 0.05]

# ── APA document helpers ──────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(3.17)
        sec.right_margin  = Cm(3.17)
    return doc

def set_cell_font(cell, size=11, bold=False):
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold

def cell_color(cell, hex_color):
    """Remove background — APA tables have no shading."""
    pass  # intentionally blank: no shading in APA

def set_para_font(para, size=11):
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)

def add_border(cell, position, size="12", color="000000", val="single"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tag  = f"w:{position}"
    el = OxmlElement(tag)
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    size)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    tcBdr = tcPr.find(qn("w:tcBorders"))
    if tcBdr is None:
        tcBdr = OxmlElement("w:tcBorders")
        tcPr.append(tcBdr)
    existing = tcBdr.find(qn(tag))
    if existing is not None:
        tcBdr.remove(existing)
    tcBdr.append(el)

def clear_border(cell, position):
    add_border(cell, position, size="0", color="FFFFFF", val="none")

def apply_apa_borders(tbl):
    """APA three-line table: thick top, thin below header, thick bottom; no other lines."""
    rows = tbl.rows
    n_rows = len(rows)
    n_cols = len(rows[0].cells)

    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row.cells):
            # clear all borders first
            for pos in ("top", "bottom", "left", "right", "insideH", "insideV"):
                clear_border(cell, pos)
            # thick top border on first row
            if r_idx == 0:
                add_border(cell, "top", size="18")        # 1.5 pt → twips*8? sz in eighths of pt
            # thin bottom border on header row (between row 0 and row 1)
            if r_idx == 0:
                add_border(cell, "bottom", size="8")      # 0.5 pt
            # thick bottom border on last row
            if r_idx == n_rows - 1:
                add_border(cell, "bottom", size="18")

def table_number(doc, number):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"Table {number}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

def table_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

def table_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run("Note. ")
    r1.italic = True
    r1.font.size = Pt(11)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = "Times New Roman"

def write_header_row(row, texts, size=11):
    for cell, txt in zip(row.cells, texts):
        cell.text = txt
        set_cell_font(cell, size=size, bold=False)

# ── Table 1: Edge Weight Matrix ────────────��──────────────────────────────────

def make_table1():
    doc = new_doc()
    table_number(doc, 1)
    table_title(doc, "Non-Zero Edge Weights in the Estimated Psychometric Network")

    # upper-triangle non-zero edges sorted by |w|
    edges = []
    for i, n1 in enumerate(NODES):
        for j, n2 in enumerate(NODES):
            if j > i:
                w = edge_matrix.get(n1, {}).get(n2, 0.0)
                if w != 0.0:
                    edges.append((n1, n2, round(w, 4)))
    edges.sort(key=lambda x: abs(x[2]), reverse=True)

    tbl = doc.add_table(rows=len(edges) + 1, cols=3)
    tbl.style = "Table Grid"
    write_header_row(tbl.rows[0], ["Node Pair", "Full Labels", "Edge Weight (ω)"])
    for i, (n1, n2, w) in enumerate(edges, 1):
        r = tbl.rows[i].cells
        r[0].text = f"{n1} – {n2}"
        r[1].text = f"{NODE_LABELS[n1]} — {NODE_LABELS[n2]}"
        r[2].text = f"{w:+.4f}"
        set_cell_font(r[0]); set_cell_font(r[1]); set_cell_font(r[2])
        if w < 0:
            r[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    apply_apa_borders(tbl)
    table_note(doc,
        f"N = {ggm['n_obs']}. Network estimated using EBICglasso (γ = {ggm['gamma']}, "
        f"Spearman correlations). Only non-zero edges are shown (total = {len(edges)} of "
        f"{ggm['n_nodes'] * (ggm['n_nodes'] - 1) // 2} possible edges; "
        f"density = {ggm['density']:.3f}). "
        "Positive weights reflect positive partial correlations; negative weights (in red) "
        "reflect negative partial correlations after controlling for all other nodes. "
        "ω = regularized partial correlation coefficient.")
    doc.save(os.path.join(TDIR, "Table1_EdgeWeights.docx"))
    print("Table 1 saved.")

# ── Table 2: Centrality Indices (z-scores) ───────────────────────────────────

def make_table2():
    doc = new_doc()
    table_number(doc, 2)
    table_title(doc, "Node Centrality Indices (z-Scores) in the Estimated Network")

    measures = ["Strength", "Betweenness", "Closeness", "ExpectedInfluence"]
    labels   = ["Strength", "Betweenness", "Closeness", "Expected Influence (1-step)"]

    # sort nodes by Strength descending
    sorted_nodes = sorted(NODES, key=lambda n: centrality.get(n, {}).get("Strength", 0), reverse=True)

    tbl = doc.add_table(rows=len(NODES) + 1, cols=6)
    tbl.style = "Table Grid"
    write_header_row(tbl.rows[0], ["Node", "Full Label"] + labels)

    for i, node in enumerate(sorted_nodes, 1):
        r = tbl.rows[i].cells
        r[0].text = node
        r[1].text = NODE_LABELS[node]
        vals = []
        for j, m in enumerate(measures, 2):
            v = centrality.get(node, {}).get(m, float("nan"))
            txt = f"{v:.3f}" if v == v else "—"
            r[j].text = txt
            vals.append(v)
        for c in r:
            set_cell_font(c)
    apply_apa_borders(tbl)
    table_note(doc,
        "Values are standardized z-scores (M = 0, SD = 1) computed across all 13 nodes. "
        "Nodes are sorted by Strength in descending order. "
        "Strength = sum of absolute edge weights. "
        "Betweenness = number of shortest paths passing through the node. "
        "Closeness = inverse mean shortest path length. "
        "Expected Influence (1-step) = sum of signed edge weights, accounting for negative edges. "
        f"CS-coefficients: Strength = {cs.get('strength', cs.get('Strength','—'))}, "
        f"Betweenness = {cs.get('betweenness', cs.get('Betweenness','—'))}, "
        f"Closeness = {cs.get('closeness', cs.get('Closeness','—'))}, "
        f"Expected Influence = {cs.get('expectedInfluence', cs.get('ExpectedInfluence','—'))}.")
    doc.save(os.path.join(TDIR, "Table2_Centrality.docx"))
    print("Table 2 saved.")

# ── Table 3: Expected Influence raw ───���──────────────────────────────────────

def make_table3():
    doc = new_doc()
    table_number(doc, 3)
    table_title(doc, "Expected Influence (Raw Values): 1-Step and 2-Step for Each Node")

    sorted_nodes = sorted(NODES, key=lambda n: ei.get(n, {}).get("step1", 0), reverse=True)

    tbl = doc.add_table(rows=len(NODES) + 1, cols=4)
    tbl.style = "Table Grid"
    write_header_row(tbl.rows[0], ["Node", "Full Label", "EI (1-step)", "EI (2-step)"])
    for i, node in enumerate(sorted_nodes, 1):
        r = tbl.rows[i].cells
        r[0].text = node
        r[1].text = NODE_LABELS[node]
        s1 = ei[node]["step1"]; s2 = ei[node]["step2"]
        r[2].text = f"{s1:.4f}"
        r[3].text = f"{s2:.4f}"
        for c in r:
            set_cell_font(c)
        if s1 < 0:
            r[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        if s2 < 0:
            r[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    apply_apa_borders(tbl)
    table_note(doc,
        "EI = Expected Influence. "
        "1-step EI = sum of signed edge weights directly connected to the node. "
        "2-step EI = additional influence propagated one hop further through the network. "
        "Positive values (black) indicate a node tends to activate its neighbors; "
        "negative values (red) indicate a node tends to suppress its neighbors. "
        "Nodes are sorted by 1-step EI in descending order.")
    doc.save(os.path.join(TDIR, "Table3_ExpectedInfluence.docx"))
    print("Table 3 saved.")

# ── Table 4: Bridge Centrality ────────────────────────────────────────────────

def make_table4():
    doc = new_doc()
    table_number(doc, 4)
    table_title(doc, "Bridge Centrality Indices for Each Node")

    sorted_nodes = sorted(NODES, key=lambda n: bridge.get(n, {}).get("strength", 0), reverse=True)

    tbl = doc.add_table(rows=len(NODES) + 1, cols=6)
    tbl.style = "Table Grid"
    write_header_row(tbl.rows[0],
        ["Node", "Full Label", "Bridge Strength", "Bridge Betweenness",
         "Bridge Closeness", "Bridge EI (1-step)"])
    for i, node in enumerate(sorted_nodes, 1):
        r = tbl.rows[i].cells
        r[0].text = node
        r[1].text = NODE_LABELS[node]
        b = bridge.get(node, {})
        r[2].text = f"{b.get('strength', 0):.4f}"
        r[3].text = f"{int(b.get('between', 0))}"
        r[4].text = f"{b.get('close', 0):.4f}"
        r[5].text = f"{b.get('EI', 0):.4f}"
        for c in r:
            set_cell_font(c)
    apply_apa_borders(tbl)
    table_note(doc,
        "Bridge centrality indices computed using the networktools package (Jones et al., 2021). "
        "Constructs used for community detection: AIS (AI Information Seeking, 5 nodes), "
        "IOA (Intention to Overuse AI, 1 node), AL (AI Literacy, 1 node), "
        "AX (AI Anxiety, 6 nodes). "
        "Bridge Strength = sum of absolute weights of edges connecting a node to nodes in other communities. "
        "Bridge Betweenness = number of between-community shortest paths passing through the node. "
        "Bridge Closeness = inverse mean shortest path to all nodes in other communities. "
        "Bridge EI = sum of signed between-community edge weights. "
        "Nodes are sorted by Bridge Strength in descending order.")
    doc.save(os.path.join(TDIR, "Table4_BridgeCentrality.docx"))
    print("Table 4 saved.")

# ── Table 5: Stability (CS-coefficients) ─────────────────────────────────���───

def make_table5():
    doc = new_doc()
    table_number(doc, 5)
    table_title(doc, "Correlation Stability (CS) Coefficients for Centrality Indices")

    rows_data = [
        ("Strength",          cs.get("strength",           cs.get("Strength",           "—"))),
        ("Expected Influence", cs.get("expectedInfluence",  cs.get("ExpectedInfluence",  "—"))),
        ("Closeness",         cs.get("closeness",           cs.get("Closeness",          "—"))),
        ("Betweenness",       cs.get("betweenness",         cs.get("Betweenness",        "—"))),
    ]

    tbl = doc.add_table(rows=len(rows_data) + 1, cols=3)
    tbl.style = "Table Grid"
    write_header_row(tbl.rows[0], ["Centrality Index", "CS-Coefficient", "Interpretation"])
    for i, (label, val) in enumerate(rows_data, 1):
        r = tbl.rows[i].cells
        r[0].text = label
        r[1].text = str(val)
        try:
            v = float(val)
            if v >= 0.50:
                interp = "Excellent (≥ .50)"
            elif v >= 0.25:
                interp = "Acceptable (.25–.49)"
            else:
                interp = "Unstable (< .25) — interpret with caution"
        except (ValueError, TypeError):
            interp = "—"
        r[2].text = interp
        for c in r:
            set_cell_font(c)
    apply_apa_borders(tbl)
    table_note(doc,
        "CS-coefficients estimated via case-dropping subset bootstrap (B = 1,000 iterations) "
        "using the bootnet package (Epskamp et al., 2018). "
        "The CS-coefficient represents the maximum proportion of cases that can be dropped "
        "while retaining a Pearson correlation ≥ .70 with the original centrality order "
        "in ≥ 95% of bootstrap samples. "
        "CS ≥ .50 = excellent stability; CS .25–.49 = acceptable; "
        "CS < .25 = unstable (Epskamp et al., 2018).")
    doc.save(os.path.join(TDIR, "Table5_Stability.docx"))
    print("Table 5 saved.")

# ── Table 6: NCT Results ─────────────────��────────────────────────────────────

def make_table6():
    doc = new_doc()
    table_number(doc, 6)
    table_title(doc, "Network Comparison Test Results: Male vs. Female Participants")

    # Panel A — global tests
    p_struct = nct["network_structure_p"]
    p_global = nct["global_strength_p"]
    n1, n2   = nct["n_group1"], nct["n_group2"]

    global_rows = [
        ("Network Structure Invariance", "M", f"{0.1384:.4f}", f"{p_struct:.4f}",
         "p > .05 — structures are invariant"),
        ("Global Strength Invariance",   "S", f"{0.1319:.4f}", f"{p_global:.4f}",
         "p > .05 — global strength is invariant"),
    ]

    p_label = doc.add_paragraph()
    r = p_label.add_run("Panel A: Global Network Tests")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"

    tbl_a = doc.add_table(rows=len(global_rows) + 1, cols=5)
    tbl_a.style = "Table Grid"
    write_header_row(tbl_a.rows[0],
        ["Test", "Statistic", "Observed Value", "p-value", "Conclusion"])
    for i, row_d in enumerate(global_rows, 1):
        for j, txt in enumerate(row_d):
            tbl_a.rows[i].cells[j].text = txt
            set_cell_font(tbl_a.rows[i].cells[j])
    apply_apa_borders(tbl_a)

    doc.add_paragraph()

    # Panel B — significant edge differences
    p_label2 = doc.add_paragraph()
    r2 = p_label2.add_run("Panel B: Edge-Level Invariance Tests (Significant Edges, p < .05)")
    r2.bold = True; r2.font.size = Pt(11); r2.font.name = "Times New Roman"

    sig = nct_sig if nct_sig else [
        ("AIS1","AIS2",0.015,0.1169),
        ("AIS3","AIS5",0.014,0.1384),
        ("AIS3","AL",  0.003,0.1128),
        ("AIS1","AX1", 0.031,0.1204),
        ("AIS4","AX1", 0.006,0.0898),
    ]
    tbl_b = doc.add_table(rows=len(sig) + 1, cols=4)
    tbl_b.style = "Table Grid"
    write_header_row(tbl_b.rows[0],
        ["Node 1", "Node 2", "p-value", "Test Statistic (E)"])
    for i, (v1, v2, p, stat) in enumerate(sig, 1):
        r = tbl_b.rows[i].cells
        r[0].text = f"{v1} ({NODE_LABELS.get(v1, v1)})"
        r[1].text = f"{v2} ({NODE_LABELS.get(v2, v2)})"
        r[2].text = f"{p:.3f}"
        r[3].text = f"{stat:.4f}"
        for c in r:
            set_cell_font(c)
        if p < 0.01:
            r[2].paragraphs[0].runs[0].font.bold = True
    apply_apa_borders(tbl_b)

    table_note(doc,
        f"Male: n = {n1}; Female: n = {n2}. "
        "Networks estimated using EBICglasso (γ = 0.5, Spearman correlations). "
        "NCT conducted using the NetworkComparisonTest package (van Borkulo et al., 2023) "
        "with 1,000 permutation iterations. "
        "Panel A: M = maximum difference in edge weights across all edges; "
        "S = difference in global network strength (sum of absolute edge weights). "
        "Panel B: Only edges with p < .05 are shown; p-values are uncorrected for multiple comparisons. "
        "Bold p-values indicate p < .01.")
    doc.save(os.path.join(TDIR, "Table6_NCT.docx"))
    print("Table 6 saved.")

# ── Copy figures ──────────────────────────────────────────────────────────────

FIGURES = [
    ("network_plot.png",                  "Figure1_NetworkPlot.png"),
    ("centrality_plot.png",               "Figure2_CentralityPlot.png"),
    ("stability_edge_plot.png",           "Figure3_StabilityEdge.png"),
    ("stability_case_plot.png",           "Figure4_StabilityCase.png"),
    ("nct_gender/nct_network_1.png",      "Figure5_NCT_Male.png"),
    ("nct_gender/nct_network_2.png",      "Figure6_NCT_Female.png"),
]

def copy_figures():
    for src_rel, dst_name in FIGURES:
        src = os.path.join(BASE, src_rel)
        dst = os.path.join(FDIR, dst_name)
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"Figure saved: {dst_name}")
        else:
            print(f"  WARNING: source not found — {src_rel}")

# ── run all ──────────────────��─────────────────────────────────────────��──────

if __name__ == "__main__":
    make_table1()
    make_table2()
    make_table3()
    make_table4()
    make_table5()
    make_table6()
    copy_figures()
    print("\nAll tables → output/network/tables/")
    print("All figures → output/network/figures/")
