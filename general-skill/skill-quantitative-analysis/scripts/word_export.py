"""
APA-format Word export for quantitative analysis results.

Produces a .docx file where every table is formatted as a 三线表:
  - Times New Roman 12 pt throughout
  - Three borders only: top of table, below header row, bottom of table
  - No vertical lines
  - Bold header row
  - Table title in italics above each table (APA convention)
  - Optional note in small print below each table
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy


# ── Border helpers ────────────────────────────────────────────────────────────

def _border_el(side: str, visible: bool, sz: int = 8) -> "OxmlElement":
    """Create a border element. Invisible borders use white so Word won't render them."""
    el = OxmlElement(f"w:{side}")
    if visible:
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    else:
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")   # white = invisible on white background
    return el


def _set_cell_borders(cell,
                      top=False, bottom=False,
                      top_sz=8, bottom_sz=8) -> None:
    """Set all six border sides for a cell. Only top/bottom can be visible."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    tcBorders.append(_border_el("top",      visible=top,    sz=top_sz))
    tcBorders.append(_border_el("bottom",   visible=bottom, sz=bottom_sz))
    tcBorders.append(_border_el("left",     visible=False))
    tcBorders.append(_border_el("right",    visible=False))
    tcBorders.append(_border_el("insideH",  visible=False))
    tcBorders.append(_border_el("insideV",  visible=False))
    tcPr.append(tcBorders)


def _clear_table_level_borders(table) -> None:
    """Set all table-level borders to white so they never bleed through."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tblBorders.append(_border_el(side, visible=False))
    tblPr.append(tblBorders)


def _apply_apa_borders(table) -> None:
    """Apply 三线表 to a table: only three black lines, everything else white."""
    _clear_table_level_borders(table)
    rows = table.rows
    if not rows:
        return
    n_rows = len(rows)
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            if r_idx == 0:
                # Header: top 1.5pt + bottom 1pt, all others white
                _set_cell_borders(cell, top=True, bottom=True,
                                   top_sz=12, bottom_sz=8)
            elif r_idx == n_rows - 1:
                # Last row: bottom 1.5pt only, all others white
                _set_cell_borders(cell, bottom=True, bottom_sz=12)
            else:
                # Middle rows: all white
                _set_cell_borders(cell)


# ── Font helper ───────────────────────────────────────────────────────────────

def _fmt_cell(cell, text: str, bold: bool = False,
              italic: bool = False, font_size: int = 12,
              align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Set cell text with Times New Roman formatting."""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)


# ── Document setup ────────────────────────────────────────────────────────────

def _setup_doc() -> Document:
    """Create a Document with Times New Roman 12 pt as the default style."""
    doc = Document()
    # Default Normal style
    normal = doc.styles["Normal"]
    normal.font.name     = "Times New Roman"
    normal.font.size     = Pt(12)
    # APA margins: 1 inch all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
    return doc


# ── Main export function ──────────────────────────────────────────────────────

def export_to_word(
    results: dict,
    filepath: str,
    title: str = "Statistical Analysis Results",
    notes: dict = None,
) -> None:
    """
    Export a dict of DataFrames to a Word document with APA 三线表 formatting.

    Parameters
    ----------
    results  : {table_title: pd.DataFrame}
               Each DataFrame becomes one APA table.
    filepath : Output .docx path (e.g. 'your-project/project-{name}/output/quantitative-analysis/results.docx')
    title    : Document-level title shown at the top.
    notes    : {table_title: note_string}  — optional APA table note below each table.

    Example
    -------
    export_to_word(
        results={
            'Table 1. Descriptive Statistics': desc_df,
            'Table 2. ANOVA Results':          anova_df,
            'Table 3. Regression Coefficients': reg_df,
        },
        filepath='your-project/project-{name}/output/quantitative-analysis/results.docx',
        notes={
            'Table 3. Regression Coefficients': 'Note. * p < .05. ** p < .01. *** p < .001.'
        }
    )
    """
    import os
    os.makedirs(os.path.dirname(filepath), exist_ok=True) if os.path.dirname(filepath) else None

    doc   = _setup_doc()
    notes = notes or {}

    # Document title
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    doc.add_paragraph()

    for table_title, df in results.items():
        if not isinstance(df, pd.DataFrame) or df.empty:
            continue

        # ── Table title (italic, left-aligned, APA style) ──────────────────
        title_para = doc.add_paragraph()
        title_run  = title_para.add_run(str(table_title))
        title_run.italic    = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)
        title_para.paragraph_format.space_after = Pt(2)

        # ── Build column list (include index if named) ──────────────────────
        include_index = bool(df.index.name)
        if include_index:
            col_headers = [str(df.index.name)] + [str(c) for c in df.columns]
        else:
            col_headers = [str(c) for c in df.columns]

        n_cols = len(col_headers)
        n_rows = len(df)

        # ── Create table ────────────────────────────────────────────────────
        tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
        # Use plain style — borders will be set manually via XML
        try:
            tbl.style = "Table Normal"
        except KeyError:
            pass  # style not present; default has no borders

        # Header row
        for c_idx, header in enumerate(col_headers):
            _fmt_cell(tbl.cell(0, c_idx), header,
                      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        for r_idx, (idx_val, row) in enumerate(df.iterrows(), start=1):
            col_offset = 0
            if include_index:
                _fmt_cell(tbl.cell(r_idx, 0), idx_val)
                col_offset = 1
            for c_idx, val in enumerate(row):
                # Format numbers: round floats, keep strings as-is
                if isinstance(val, float):
                    display = f"{val:.3f}" if abs(val) >= 0.001 else f"{val:.2e}"
                else:
                    display = str(val) if val is not None else ""
                _fmt_cell(tbl.cell(r_idx, c_idx + col_offset), display,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Apply 三线表 borders
        _apply_apa_borders(tbl)

        # ── Optional APA note ───────────────────────────────────────────────
        note_text = notes.get(table_title, "")
        if note_text:
            note_para = doc.add_paragraph()
            note_run  = note_para.add_run(note_text)
            note_run.italic    = True
            note_run.font.name = "Times New Roman"
            note_run.font.size = Pt(10)
            note_para.paragraph_format.space_before = Pt(2)

        doc.add_paragraph()   # space between tables

    doc.save(filepath)
    print(f"Word document saved: {filepath}")
