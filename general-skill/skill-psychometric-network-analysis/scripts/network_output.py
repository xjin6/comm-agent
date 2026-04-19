"""
network_output.py — APA-style output for Psychometric Network Analysis skill
Generates Word tables, Excel workbook, and bilingual write-up documents.

Usage:
    python network_output.py <output_dir> <tables> [--lang en|cn|both]

    tables: comma-separated subset of:
        descriptives, edge_weights, centrality, stability, nct, clpna
"""

import argparse
import json
import os
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# ── APA helpers ──────────────────────────────────��────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE = 12


def new_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
    return doc


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "12")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_apa_table(doc: Document, title: str, note: str,
                  headers: list, rows: list) -> None:
    """Add an APA three-line table (top line, header line, bottom line)."""
    doc.add_paragraph()
    p = doc.add_paragraph(title)
    p.runs[0].bold = True
    p.runs[0].font.name = FONT_NAME
    p.runs[0].font.size = Pt(FONT_SIZE)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val) if val is not None else ""
            cells[c_idx].paragraphs[0].runs[0].font.name = FONT_NAME
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(FONT_SIZE)

    # Three-line borders: top of header, bottom of header, bottom of last row
    n_cols = len(headers)
    for c in tbl.rows[0].cells:
        _set_cell_border(c, top="single", bottom="single",
                         left="none", right="none")
    for c in tbl.rows[-1].cells:
        _set_cell_border(c, bottom="single")

    # Remove all other borders
    for row in tbl.rows[1:]:
        for cell in row.cells:
            _set_cell_border(cell, left="none", right="none")

    if note:
        p_note = doc.add_paragraph()
        p_note.add_run("Note. ").italic = True
        p_note.add_run(note)
        for run in p_note.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)


# ── Table generators ──────────────────────────────────────────────────────────

def table_descriptives(doc: Document, output_dir: Path) -> None:
    csv_path = output_dir / "edge_weights.csv"
    if not csv_path.exists():
        doc.add_paragraph("[edge_weights.csv not found — run GGM first]")
        return

    df  = pd.read_csv(csv_path, index_col=0)
    nodes = list(df.columns)

    desc_rows = []
    for node in nodes:
        col_data = df[node].dropna()
        desc_rows.append([node,
                          f"{col_data.mean():.2f}",
                          f"{col_data.std():.2f}",
                          f"{col_data.min():.2f}",
                          f"{col_data.max():.2f}"])

    add_apa_table(
        doc,
        title="Table 1\nDescriptive Statistics",
        note="Values derived from partial correlation edge weight matrix.",
        headers=["Node", "M", "SD", "Min", "Max"],
        rows=desc_rows
    )


def table_edge_weights(doc: Document, output_dir: Path) -> None:
    csv_path = output_dir / "edge_weights.csv"
    if not csv_path.exists():
        doc.add_paragraph("[edge_weights.csv not found]")
        return

    df    = pd.read_csv(csv_path, index_col=0)
    nodes = list(df.columns)
    rows  = []
    for node in nodes:
        row = [node] + [f"{df.loc[node, col]:.3f}" if node != col else "—"
                        for col in nodes]
        rows.append(row)

    add_apa_table(
        doc,
        title="Table 2\nEBICglasso Partial Correlation Edge Weights",
        note=(
            "Edge weights are regularized partial correlations estimated via "
            "EBICglasso (γ = 0.5). Positive edges are shown in blue; "
            "negative edges in red in the network plot."
        ),
        headers=[""] + nodes,
        rows=rows
    )


def table_centrality(doc: Document, output_dir: Path) -> None:
    ci_path = output_dir / "centrality_indices.csv"
    ei_path = output_dir / "expected_influence.csv"
    if not ci_path.exists():
        doc.add_paragraph("[centrality_indices.csv not found]")
        return

    ci = pd.read_csv(ci_path)
    # centralityTable returns long format: node, measure, value
    if "measure" in ci.columns and "value" in ci.columns:
        ci_wide = ci.pivot(index="node", columns="measure", values="value").reset_index()
    else:
        ci_wide = ci

    # Merge EI
    if ei_path.exists():
        ei = pd.read_csv(ei_path)
        ci_wide = ci_wide.merge(ei, left_on="node", right_on="node", how="left")

    headers = list(ci_wide.columns)
    rows    = ci_wide.round(3).values.tolist()

    # CS note
    cs_path = output_dir / "cs_coefficients.csv"
    note = (
        "Centrality indices estimated from EBICglasso network. "
        "EI = Expected Influence."
    )
    if cs_path.exists():
        cs = pd.read_csv(cs_path)
        cs_str = "; ".join(
            f"{row['statistic']} CS = {row['CS_coef']:.2f}"
            for _, row in cs.iterrows()
        )
        note += f" Case-dropping bootstrap stability: {cs_str}. "
        note += "CS ≥ 0.50 indicates high stability; CS ≥ 0.25 is acceptable."

    add_apa_table(
        doc,
        title="Table 3\nCentrality Indices",
        note=note,
        headers=headers,
        rows=rows
    )


def table_stability(doc: Document, output_dir: Path) -> None:
    cs_path = output_dir / "cs_coefficients.csv"
    if not cs_path.exists():
        doc.add_paragraph("[cs_coefficients.csv not found — run stability first]")
        return

    cs = pd.read_csv(cs_path)
    rows = cs.round(3).values.tolist()

    add_apa_table(
        doc,
        title="Table 4\nBootstrap Stability Results (CS-Coefficients)",
        note=(
            "CS-coefficients estimated via case-dropping subset bootstrap "
            "(B = 1,000). Values ≥ 0.50 indicate high stability; "
            "values ≥ 0.25 are acceptable (Epskamp et al., 2018)."
        ),
        headers=["Centrality Index", "CS-Coefficient"],
        rows=rows
    )


def table_nct(doc: Document, output_dir: Path) -> None:
    json_path = output_dir / "nct_summary.json"
    txt_path  = output_dir / "nct_summary.txt"
    if not json_path.exists():
        doc.add_paragraph("[nct_summary.json not found — run NCT first]")
        return

    with open(json_path) as f:
        nct = json.load(f)

    rows = [
        ["Global strength invariance test", "M",
         nct.get("global_strength_p", "—")],
        ["Network structure invariance test", "M",
         nct.get("network_structure_p", "—")],
    ]

    add_apa_table(
        doc,
        title="Table 5\nNetwork Comparison Test Results",
        note=(
            f"Group 1: {nct.get('group1','—')} (n = {nct.get('n_group1','—')}); "
            f"Group 2: {nct.get('group2','—')} (n = {nct.get('n_group2','—')}). "
            "NCT performed with 1,000 permutations "
            "(van Borkulo et al., 2022). M = test statistic."
        ),
        headers=["Test", "Statistic", "p"],
        rows=rows
    )


def table_clpna(doc: Document, output_dir: Path) -> None:
    temp_path = output_dir / "mlvar_temporal_matrix.csv"
    pvar_path = output_dir / "panelvar_coefficients.csv"

    if temp_path.exists():
        df   = pd.read_csv(temp_path, index_col=0)
        rows = [[idx] + [f"{v:.3f}" for v in row]
                for idx, row in zip(df.index, df.values)]
        add_apa_table(
            doc,
            title="Table 6\nCross-lagged Temporal Paths (mlVAR)",
            note=(
                "Standardized temporal (cross-lagged) coefficients from "
                "multilevel VAR model. Rows = outcomes; columns = predictors "
                "(lagged). Estimated via mlVAR (Epskamp et al., 2018)."
            ),
            headers=[""] + list(df.columns),
            rows=rows
        )
    elif pvar_path.exists():
        df   = pd.read_csv(pvar_path, index_col=0)
        rows = [[idx] + [f"{v:.3f}" if pd.notna(v) else "—" for v in row]
                for idx, row in zip(df.index, df.values)]
        add_apa_table(
            doc,
            title="Table 6\nCross-lagged Paths (panelvar GMM)",
            note=(
                "GMM coefficient estimates from panelvar two-step first-difference "
                "estimator. Rows = lagged predictors; columns = outcomes."
            ),
            headers=[""] + list(df.columns),
            rows=rows
        )
    else:
        doc.add_paragraph("[No CLPNA output found — run CLPNA first]")


# ── Write-up generators ───────────────────────────────────────────────────────

EN_WRITEUP = """\
Network Analysis Results

Cross-Sectional Network

A Gaussian Graphical Model (GGM) was estimated using EBICglasso regularization \
(γ = 0.5) via the bootnet package in R (Epskamp & Fried, 2018). \
The regularized partial correlation network is depicted in Figure 1 \
(positive edges = blue; negative edges = red). \
Edge weights and centrality indices (Strength, Betweenness, Closeness, \
Expected Influence) are reported in Tables 2 and 3.

Stability of centrality indices was examined using a case-dropping subset \
bootstrap procedure (B = 1,000; Epskamp et al., 2018). \
CS-coefficients are reported in Table 4. \
Values ≥ 0.50 indicate high stability; values ≥ 0.25 are acceptable.

References

Epskamp, S., & Fried, E. I. (2018). A tutorial on regularized partial \
correlation networks. Psychological Methods, 23(4), 617–634. \
https://doi.org/10.1037/met0000167

van Borkulo, C. D., Borsboom, D., Epskamp, S., Blanken, T. F., Boschloo, L., \
Schoevers, R. A., & Waldorp, L. J. (2022). Comparing network structures on \
three aspects: A permutation test. Psychological Methods, 27(6), 1033–1043. \
https://doi.org/10.1037/met0000476
"""

CN_WRITEUP = """\
网络分析结果

横截面网络

使用 R 语言 bootnet 包，通过 EBICglasso 正则化（γ = 0.5）估计高斯图模型（GGM）\
（Epskamp & Fried, 2018）。图 1 展示了正则化偏相关网络图（蓝色为正向连边，红色为负向连边）。\
边权重及各中心性指标（强度、中介、接近、期望影响力）详见表 2 和表 3。

采用逐步剔除被试的自助法（B = 1,000）检验中心性指标的稳定性\
（Epskamp 等, 2018），CS 系数见表 4。\
CS ≥ 0.50 表明稳定性高，CS ≥ 0.25 为可接受水平。

参考文献

Epskamp, S., & Fried, E. I. (2018). A tutorial on regularized partial \
correlation networks. Psychological Methods, 23(4), 617–634. \
https://doi.org/10.1037/met0000167
"""


def write_writeup(output_dir: Path, lang: str) -> None:
    if lang in ("en", "both"):
        doc = new_doc()
        for line in EN_WRITEUP.strip().split("\n"):
            doc.add_paragraph(line)
        doc.save(output_dir / "NetworkAnalysis_EN.docx")
        print("Saved NetworkAnalysis_EN.docx")

    if lang in ("cn", "both"):
        doc = new_doc()
        for line in CN_WRITEUP.strip().split("\n"):
            doc.add_paragraph(line)
        doc.save(output_dir / "NetworkAnalysis_CN.docx")
        print("Saved NetworkAnalysis_CN.docx")


# ── Excel workbook ────────────────────────────────────────────────────────────

def write_excel(output_dir: Path, tables: list) -> None:
    wb   = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    thin = Side(style="thin")
    thick = Side(style="medium")

    def add_sheet(name: str, csv_path: Path) -> None:
        if not csv_path.exists():
            return
        df = pd.read_csv(csv_path)
        ws = wb.create_sheet(title=name[:31])
        headers = list(df.columns)

        # Header row
        for c_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=c_idx, value=h)
            cell.font      = Font(name="Times New Roman", size=12, bold=True)
            cell.alignment = Alignment(horizontal="center")
            cell.border    = Border(top=thick, bottom=thick)

        # Data rows
        for r_idx, row in enumerate(df.itertuples(index=False), start=2):
            for c_idx, val in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                cell.font      = Font(name="Times New Roman", size=12)
                cell.alignment = Alignment(horizontal="center")

        # Bottom border on last row
        last_row = ws.max_row
        for c_idx in range(1, len(headers) + 1):
            ws.cell(row=last_row, column=c_idx).border = Border(bottom=thick)

    csv_map = {
        "EdgeWeights":  output_dir / "edge_weights.csv",
        "Centrality":   output_dir / "centrality_indices.csv",
        "ExpInfl":      output_dir / "expected_influence.csv",
        "BridgeCent":   output_dir / "bridge_centrality.csv",
        "Stability_CS": output_dir / "cs_coefficients.csv",
        "NCT":          output_dir / "nct_summary.json",
        "TempMatrix":   output_dir / "mlvar_temporal_matrix.csv",
        "ContempMatrix":output_dir / "mlvar_contemporaneous_matrix.csv",
        "PanelVAR":     output_dir / "panelvar_coefficients.csv",
    }

    for sheet_name, csv_path in csv_map.items():
        if csv_path.suffix == ".json":
            continue  # skip JSON for Excel
        add_sheet(sheet_name, csv_path)

    if len(wb.sheetnames) == 0:
        wb.create_sheet("Empty")

    out_path = output_dir / "network_results.xlsx"
    wb.save(out_path)
    print(f"Saved {out_path}")


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output_dir")
    parser.add_argument("tables", help="comma-separated: descriptives,edge_weights,"
                                       "centrality,stability,nct,clpna,all")
    parser.add_argument("--lang", default="both", choices=["en", "cn", "both"])
    parser.add_argument("--writeup", action="store_true",
                        help="Also generate bilingual write-up documents")
    parser.add_argument("--excel", action="store_true",
                        help="Also generate Excel workbook")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    table_set = {t.strip() for t in args.tables.split(",")}
    if "all" in table_set:
        table_set = {"descriptives", "edge_weights", "centrality",
                     "stability", "nct", "clpna"}

    doc = new_doc()

    dispatch = {
        "descriptives": table_descriptives,
        "edge_weights": table_edge_weights,
        "centrality":   table_centrality,
        "stability":    table_stability,
        "nct":          table_nct,
        "clpna":        table_clpna,
    }

    for key in ["descriptives", "edge_weights", "centrality",
                "stability", "nct", "clpna"]:
        if key in table_set:
            dispatch[key](doc, output_dir)

    out_docx = output_dir / "NetworkAnalysis_Tables.docx"
    doc.save(out_docx)
    print(f"Saved {out_docx}")

    if args.writeup:
        write_writeup(output_dir, args.lang)

    if args.excel:
        write_excel(output_dir, table_set)


if __name__ == "__main__":
    main()
