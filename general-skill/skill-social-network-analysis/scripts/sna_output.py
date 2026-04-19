#!/usr/bin/env python3
"""
sna_output.py — Social Network Analysis: APA-style Word tables, centrality
                 CSV export, Gephi .gexf export, and combined report generation.
Called by the SNA skill after sna_analysis.R has completed.
"""

import os
import sys
import csv
import argparse
import networkx as nx
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Helpers: APA three-line table formatting
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        val = kwargs.get(edge, {})
        tag.set(qn("w:val"),   val.get("val",   "none"))
        tag.set(qn("w:sz"),    str(val.get("sz",  0)))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), val.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _set_row_border(row, position: str, sz_eighths: int):
    """Set a horizontal border on all cells in a row."""
    for cell in row.cells:
        border_kwargs = {
            position: {"val": "single", "sz": sz_eighths, "color": "000000"}
        }
        _set_cell_border(cell, **border_kwargs)


def apa_table(doc: Document, title: str, headers: list, rows: list) -> None:
    """Add an APA three-line table to *doc*."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(title).bold = True

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None else ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # APA borders: top 1.5 pt (12 eighths), header-bottom 0.5 pt (4 eighths),
    #              table-bottom 1.5 pt (12 eighths); no verticals
    _set_row_border(tbl.rows[0], "top",    12)
    _set_row_border(tbl.rows[0], "bottom",  4)
    _set_row_border(tbl.rows[-1], "bottom", 12)


# ---------------------------------------------------------------------------
# Table 1 — Network Descriptives
# ---------------------------------------------------------------------------

def make_table1_descriptives(output_dir: Path, doc: Document) -> None:
    csv_path = output_dir / "descriptives.csv"
    if not csv_path.exists():
        return

    df = pd.read_csv(csv_path)
    headers = ["Metric", "Value"]
    rows = [(row["metric"], row["value"]) for _, row in df.iterrows()]
    apa_table(doc, "Table 1\nNetwork-Level Descriptive Statistics", headers, rows)


# ---------------------------------------------------------------------------
# Table 2 — Centrality Indices
# ---------------------------------------------------------------------------

def make_table2_centrality(output_dir: Path, doc: Document) -> None:
    csv_path = output_dir / "centrality_indices.csv"
    if not csv_path.exists():
        return

    df = pd.read_csv(csv_path)
    df_sorted = df.sort_values(df.columns[1], ascending=False)  # sort by first centrality col

    headers = list(df_sorted.columns)
    rows = [
        [f"{v:.4f}" if isinstance(v, float) else str(v) for v in row]
        for row in df_sorted.itertuples(index=False)
    ]
    apa_table(doc, "Table 2\nCentrality Indices", headers, rows)


# ---------------------------------------------------------------------------
# Table 3 — Key Nodes
# ---------------------------------------------------------------------------

def make_table3_key_nodes(output_dir: Path, doc: Document) -> None:
    csv_path = output_dir / "key_nodes.csv"
    if not csv_path.exists():
        return

    df = pd.read_csv(csv_path)
    headers = list(df.columns)
    rows = [
        [f"{v:.4f}" if isinstance(v, float) else str(v) for v in row]
        for row in df.itertuples(index=False)
    ]
    apa_table(doc, "Table 3\nKey Nodes — Top N by Composite Centrality Score", headers, rows)


# ---------------------------------------------------------------------------
# Table 4 — Community Detection
# ---------------------------------------------------------------------------

def make_table4_community(output_dir: Path, doc: Document) -> None:
    mem_path  = output_dir / "community_membership.csv"
    stat_path = output_dir / "community_stats.txt"
    if not mem_path.exists():
        return

    mem_df = pd.read_csv(mem_path)
    comm_sizes = mem_df.groupby("community").size().reset_index(name="size")
    comm_sizes = comm_sizes.sort_values("size", ascending=False)

    # Read modularity if available
    modularity_str = "—"
    if stat_path.exists():
        for line in stat_path.read_text().splitlines():
            if line.startswith("modularity"):
                modularity_str = line.split(",")[1].strip()

    headers = ["Community", "Size (Nodes)", "% of Network"]
    total   = mem_df.shape[0]
    rows = [
        (int(row["community"]),
         int(row["size"]),
         f"{100 * row['size'] / total:.1f}%")
        for _, row in comm_sizes.iterrows()
    ]

    apa_table(doc,
              f"Table 4\nCommunity Detection Results  (Modularity = {modularity_str})",
              headers, rows)


# ---------------------------------------------------------------------------
# Table 5 — Diffusion Analysis
# ---------------------------------------------------------------------------

def make_table5_diffusion(output_dir: Path, doc: Document) -> None:
    csv_path = output_dir / "diffusion_analysis.csv"
    if not csv_path.exists():
        return

    df = pd.read_csv(csv_path)
    headers = list(df.columns)
    rows = [
        [f"{v:.4f}" if isinstance(v, float) else str(v) for v in row]
        for row in df.itertuples(index=False)
    ]
    apa_table(doc, "Table 5\nInformation Diffusion / Cascade Analysis", headers, rows)


# ---------------------------------------------------------------------------
# Insert figures
# ---------------------------------------------------------------------------

FIGURES = [
    ("figures/Figure1_NetworkPlot.png",        "Figure 1. Social network plot. Node size proportional to PageRank; colour indicates community membership."),
    ("figures/Figure2_CentralityPlot.png",     "Figure 2. Top-20 nodes by centrality measures."),
    ("figures/Figure3_DegreeDistribution.png", "Figure 3. Degree distribution of the network."),
    ("figures/Figure4_CommunityStructure.png", "Figure 4. Community structure highlighting."),
    ("figures/Figure5_DiffusionCascade.png",   "Figure 5. Information diffusion cascade tree."),
]


def add_figures(output_dir: Path, doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for rel_path, caption in FIGURES:
        fig_path = output_dir / rel_path
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Cm(14))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate SNA APA output")
    parser.add_argument("--output-dir", required=True,
                        help="Path to output/network/sna/ directory")
    parser.add_argument("--graphml",    required=True,
                        help="Path to network GraphML file for Gephi export")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    tables_dir = output_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "figures").mkdir(parents=True, exist_ok=True)

    # --- Individual tables ---
    table_funcs = [
        ("Table1_Descriptives.docx",  make_table1_descriptives),
        ("Table2_Centrality.docx",    make_table2_centrality),
        ("Table3_KeyNodes.docx",      make_table3_key_nodes),
        ("Table4_Community.docx",     make_table4_community),
        ("Table5_Diffusion.docx",     make_table5_diffusion),
    ]

    for fname, fn in table_funcs:
        doc = Document()
        fn(output_dir, doc)
        doc.save(tables_dir / fname)
        print(f"Saved {fname}")

    # --- Combined report ---
    report = Document()
    report.add_heading("Social Network Analysis Report", level=0)
    report.add_paragraph()

    for _, fn in table_funcs:
        fn(output_dir, report)

    add_figures(output_dir, report)
    report.save(output_dir / "sna_report.docx")
    print("Saved sna_report.docx")

    # --- Gephi .gexf export ---
    if args.graphml and Path(args.graphml).exists():
        G = nx.read_graphml(args.graphml)
        gexf_path = output_dir / "network_gephi.gexf"
        nx.write_gexf(G, str(gexf_path))
        print(f"Saved network_gephi.gexf ({G.number_of_nodes()} nodes, {G.number_of_edges()} edges)")


if __name__ == "__main__":
    main()
